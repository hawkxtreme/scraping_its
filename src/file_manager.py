import os
import shutil
import json
import re
import time
from urllib.parse import urlparse
import markdownify

try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from . import config
except ImportError:
    import config

def setup_output_directories(formats, update_mode=False, log_func=None):
    """Cleans and creates the necessary output directories based on specified formats."""
    # Create base output directory if it doesn't exist
    base_dir = config.get_output_dir()
    os.makedirs(base_dir, exist_ok=True)
    
    if log_func:
        log_func.logger.log_step("Setting up output directories", {
            "formats": formats,
            "update_mode": update_mode,
            "base_directory": base_dir
        })

    # Clean and create directories for specified formats
    format_map = {
        'json': config.get_json_dir(),
        'pdf': config.get_pdf_dir(),
        'txt': config.get_txt_dir(),
        'markdown': config.get_markdown_dir(),
        'docx': config.get_docx_dir(),
    }

    for fmt in formats:
        dir_path = format_map.get(fmt)
        if dir_path:
            # In update mode, don't delete existing directories
            if not update_mode and os.path.exists(dir_path):
                if log_func:
                    log_func.logger.debug(f"Removing existing directory for format {fmt}: {dir_path}")
                shutil.rmtree(dir_path)
            
            os.makedirs(dir_path, exist_ok=True)
            if log_func:
                log_func.logger.debug(f"Ensured directory exists for format {fmt}: {dir_path}")

    # Ensure tmp_index directory exists
    tmp_dir = config.get_tmp_index_dir()
    os.makedirs(tmp_dir, exist_ok=True)
    if log_func:
        log_func.logger.debug(f"Ensured temp index directory exists: {tmp_dir}")

def save_article_json(article_data):
    """Save article in JSON format."""
    json_dir = config.get_json_dir()
    if not os.path.exists(json_dir):
        os.makedirs(json_dir)

    # Get just the path component of the URL for the filename
    url_path = urlparse(article_data['url']).path
    safe_filename = re.sub(r'[^\w\-_.]', '_', url_path.strip('/'))
    
    json_file = os.path.join(json_dir, f"{safe_filename}.json")
    with open(json_file, 'w', encoding='utf-8') as f:
        json.dump(article_data, f, ensure_ascii=False, indent=2)

def _save_article_json_and_txt_legacy(content, article_info):
    """Save article content in both JSON and TXT formats."""
    article_data = {
        'url': article_info['url'],
        'title': article_info['title'],
        'content': content
    }
    save_article_json(article_data)

    # TXT format
    txt_dir = config.get_txt_dir()
    if not os.path.exists(txt_dir):
        os.makedirs(txt_dir)
    
    # Get just the path component of the URL for the filename
    url_path = urlparse(article_info['url']).path
    safe_filename = re.sub(r'[^\w\-_.]', '_', url_path.strip('/'))
    
    txt_file = os.path.join(txt_dir, f"{safe_filename}.txt")
    with open(txt_file, 'w', encoding='utf-8') as f:
        f.write(f"Title: {article_info['title']}\n")
        f.write(f"URL: {article_info['url']}\n\n")
        f.write(content)
def save_article_content(filename_base, formats, soup, article_info, log_func=None):
    """
    Сохраняет статью в указанных форматах (txt, json, markdown, pdf) по имени файла.
    - filename_base: базовое имя файла (без расширения)
    - formats: список форматов (например, ['txt', 'json', 'markdown'])
    - soup: BeautifulSoup-объект с содержимым статьи
    - article_info: словарь с метаданными статьи
    - log_func: функция логирования (опционально)
    """
    save_start = time.time()
    article_title = article_info.get('title', 'Unknown Article')
    
    if log_func:
        log_func.logger.debug(f"Saving article content for: {article_title}")
    
    text_content = soup.get_text(separator='\n', strip=True)
    article_data = {
        'url': article_info['url'],
        'title': article_info['title'],
        'content': text_content
    }

    saved_formats = []

    # JSON
    if 'json' in formats:
        try:
            json_file = os.path.join(config.get_json_dir(), f"{filename_base}.json")
            with open(json_file, 'w', encoding='utf-8') as f:
                json.dump(article_data, f, ensure_ascii=False, indent=2)
            saved_formats.append('json')
            if log_func:
                log_func.logger.debug(f"Saved JSON file: {json_file}")
        except Exception as e:
            if log_func:
                log_func.logger.error(f"Failed to save JSON file for {article_title}", e)

    # TXT
    if 'txt' in formats:
        try:
            txt_file = os.path.join(config.get_txt_dir(), f"{filename_base}.txt")
            with open(txt_file, 'w', encoding='utf-8') as f:
                f.write(f"Title: {article_info['title']}\n")
                f.write(f"URL: {article_info['url']}\n\n")
                f.write(text_content)
            saved_formats.append('txt')
            if log_func:
                log_func.logger.debug(f"Saved TXT file: {txt_file}")
        except Exception as e:
            if log_func:
                log_func.logger.error(f"Failed to save TXT file for {article_title}", e)

    # Markdown
    if 'markdown' in formats:
        try:
            md_file = os.path.join(config.get_markdown_dir(), f"{filename_base}.md")
            md_content = markdownify.markdownify(str(soup), heading_style="ATX")
            with open(md_file, 'w', encoding='utf-8') as f:
                f.write(md_content)
            saved_formats.append('markdown')
            if log_func:
                log_func.logger.debug(f"Saved Markdown file: {md_file}")
        except Exception as e:
            if log_func:
                log_func.logger.error(f"Failed to save Markdown file for {article_title}", e)

    # DOCX
    if 'docx' in formats:
        if not DOCX_AVAILABLE:
            if log_func:
                log_func.logger.warning("python-docx library not available, skipping DOCX format")
        else:
            try:
                docx_file = os.path.join(config.get_docx_dir(), f"{filename_base}.docx")
                _save_as_docx(soup, article_info, docx_file, log_func)
                saved_formats.append('docx')
                if log_func:
                    log_func.logger.debug(f"Saved DOCX file: {docx_file}")
            except Exception as e:
                if log_func:
                    log_func.logger.error(f"Failed to save DOCX file for {article_title}", e)

    # PDF сохраняется отдельно через scraper._save_as_pdf
    
    save_duration = time.time() - save_start
    if log_func:
        log_func.logger.log_performance(f"Content saving for {article_title}", save_duration, {
            "formats": saved_formats
        })

def save_hierarchical_index(toc_tree, log_func=None):
    """Saves the hierarchical TOC tree to a JSON file."""
    index_file = os.path.join(config.get_tmp_index_dir(), "_toc_tree.json")
    
    try:
        with open(index_file, 'w', encoding='utf-8') as f:
            json.dump(toc_tree, f, ensure_ascii=False, indent=2)
        
        if log_func:
            log_func.logger.info(f"Saved hierarchical index with {len(toc_tree)} top-level items")
    except Exception as e:
        if log_func:
            log_func.logger.error("Failed to save hierarchical index", e)
        raise

def create_toc_and_meta(scraped_articles, all_articles, formats, log_func=None):
    """Creates the _toc.md and _meta.json files."""
    # We need the original toc_tree to generate the markdown toc
    index_file = os.path.join(config.get_tmp_index_dir(), "_toc_tree.json")
    if not os.path.exists(index_file):
        if log_func:
            log_func.logger.error("TOC tree file not found, cannot create TOC and meta files")
        return
    
    try:
        with open(index_file, "r", encoding="utf-8") as f:
            toc_tree = json.load(f)

        create_markdown_toc(toc_tree, all_articles, formats, log_func)
        
        # Ensure each article has a content_hash for update tracking
        for article in all_articles:
            if 'content_hash' not in article:
                article['content_hash'] = None
        
        create_meta_json(scraped_articles, all_articles, formats, log_func)
        
        if log_func:
            log_func.logger.info(f"Created TOC and metadata files for {len(scraped_articles)} scraped articles, {len(all_articles)} total articles")
    except Exception as e:
        if log_func:
            log_func.logger.error("Failed to create TOC and metadata files", e)
        raise


def create_markdown_toc(toc_tree, articles, formats, log_func=None):
    """Generates the _toc.md file with all existing files."""
    url_to_filename = {article["url"]: article["filename_base"] for article in articles if "filename_base" in article}
    
    # Scan all format directories to find existing files
    format_dirs = {
        'json': config.get_json_dir(),
        'pdf': config.get_pdf_dir(),
        'txt': config.get_txt_dir(),
        'markdown': config.get_markdown_dir(),
        'docx': config.get_docx_dir(),
    }
    
    # Create a mapping from filename_base to existing formats
    filename_to_formats = {}
    
    for format_name, dir_path in format_dirs.items():
        if os.path.exists(dir_path):
            for file in os.listdir(dir_path):
                if file.endswith(f".{format_name}"):
                    filename_base = os.path.splitext(file)[0]
                    if filename_base not in filename_to_formats:
                        filename_to_formats[filename_base] = []
                    filename_to_formats[filename_base].append(format_name)
    
    # Create a reverse mapping from filename_base to article info
    filename_to_article = {}
    for article in articles:
        if "filename_base" in article:
            filename_to_article[article["filename_base"]] = article

    try:
        toc_file = os.path.join(config.get_output_dir(), "_toc.md")
        with open(toc_file, "w", encoding="utf-8") as f:
            f.write("# Оглавление\n\n")
            
            def write_nodes(nodes, indent_level=0):
                for node in nodes:
                    filename_base = url_to_filename.get(node["url"])
                    
                    if filename_base:
                        f.write("    " * indent_level + f"*   **{node['title']}**\n")
                        
                        # Get all available formats for this file
                        available_formats = filename_to_formats.get(filename_base, [])
                        
                        # Add links for all available formats
                        for format_name in sorted(available_formats):
                            f.write("    " * (indent_level + 1) + f"*   [{format_name.upper()}](./{format_name}/{filename_base}.{format_name})\n")
                    else:
                        f.write("    " * indent_level + f"*   **{node['title']}**\n")

                    if node["children"]:
                        write_nodes(node["children"], indent_level + 1)
            
            # First, write the TOC based on the tree structure
            write_nodes(toc_tree)
            
            # Then, add a section for files that exist but aren't in the TOC
            f.write("\n## Дополнительные файлы\n\n")
            for filename_base, formats_list in sorted(filename_to_formats.items()):
                if filename_base not in url_to_filename.values():
                    article_info = filename_to_article.get(filename_base)
                    title = article_info.get('title', filename_base) if article_info else filename_base
                    
                    f.write(f"*   **{title}**\n")
                    for format_name in sorted(formats_list):
                        f.write(f"    *   [{format_name.upper()}](./{format_name}/{filename_base}.{format_name})\n")
        
        if log_func:
            log_func.logger.info(f"Created markdown TOC file: {toc_file}")
            log_func.logger.info(f"TOC includes {len(filename_to_formats)} files")
    except Exception as e:
        if log_func:
            log_func.logger.error("Failed to create markdown TOC file", e)
        raise

def create_meta_json(scraped_articles, all_articles, formats, log_func=None):
    """Generates the _meta.json file."""
    try:
        meta_file = os.path.join(config.get_output_dir(), "_meta.json")

        # Create a URL-based lookup for scraped articles to easily access their content hashes
        scraped_lookup = {
            article.get("url"): article.get("content_hash")
            for article in scraped_articles
            if article.get("url") and article.get("content_hash")
        }

        # Create a URL-based lookup from the existing _meta.json file to preserve hashes
        # of articles that were not scraped in this run.
        existing_meta = load_existing_meta_data()
        existing_lookup = {
            article.get("url"): article.get("content_hash")
            for article in existing_meta
            if article.get("url") and article.get("content_hash")
        }

        # Iterate through all articles that should be in the final meta file
        for article in all_articles:
            url = article.get("url")
            if not url:
                continue

            # If the article was scraped in this run, use its new hash
            if url in scraped_lookup:
                article["content_hash"] = scraped_lookup[url]
            # Otherwise, if it exists in the old meta file, preserve its old hash
            elif url in existing_lookup:
                article["content_hash"] = existing_lookup[url]
            # If it's a new article that wasn't scraped (e.g. --no-scrape), it won't have a hash
            # The 'create_toc_and_meta' function already initializes it to None, so no action needed here.

        # Save the updated list of all articles to the meta file
        with open(meta_file, "w", encoding="utf-8") as f:
            json.dump(all_articles, f, ensure_ascii=False, indent=4)

        if log_func:
            log_func.logger.info(f"Created metadata JSON file: {meta_file} with {len(all_articles)} total articles")
            updated_count = len(scraped_lookup)
            preserved_count = len(all_articles) - updated_count
            log_func.logger.info(f"Updated {updated_count} articles, preserved hashes for {preserved_count} articles.")

    except Exception as e:
        if log_func:
            log_func.logger.error("Failed to create metadata JSON file", e)
        raise

def load_index_data():
    """Loads the hierarchical index, flattens it, and returns a list of articles to scrape."""
    index_file = os.path.join(config.get_tmp_index_dir(), "_toc_tree.json")
    if not os.path.exists(index_file):
        return []
        
    with open(index_file, "r", encoding="utf-8") as f:
        toc_tree = json.load(f)

    flat_list = []
    # Helper to sanitize titles for use in filenames
    def sanitize_title(title):
        # Remove invalid characters for filenames
        sanitized = re.sub(r'[<>:"/\\|?*]', '', title)
        # Replace long sequences of whitespace with a single underscore
        sanitized = re.sub(r'\s+', '_', sanitized)
        # Truncate to a reasonable length
        return sanitized[:100]

    def traverse(nodes, breadcrumbs, counter):
        for node in nodes:
            # Only process nodes that have a URL (i.e., are articles)
            if "url" in node and node["url"]:
                # Format index with zero padding
                index_str = str(counter).zfill(4) # e.g., 0001, 0002
                sanitized_title = sanitize_title(node["title"])
                filename_base = f"{index_str}_{sanitized_title}"

                article_data = {
                    "index": counter,
                    "filename_base": filename_base,
                    "title": node["title"],
                    "url": node["url"],
                    "breadcrumb": breadcrumbs + [node["title"]],
                }
                flat_list.append(article_data)
                counter += 1
            
            # Recursively process children, carrying the counter forward
            if "children" in node and node["children"]:
                counter = traverse(node["children"], breadcrumbs + [node["title"]], counter)
        return counter

    traverse(toc_tree, [], 1) # Start indexing from 1
    return flat_list

def get_index_path():
    """Returns the path to the temporary index directory."""
    return config.get_tmp_index_dir()

def cleanup_temp_files(log_func=None):
    """Removes temporary directories."""
    tmp_index_dir = config.get_tmp_index_dir()
    if os.path.exists(tmp_index_dir):
        try:
            shutil.rmtree(tmp_index_dir)
            if log_func:
                log_func.logger.debug(f"Cleaned up temporary directory: {tmp_index_dir}")
        except Exception as e:
            if log_func:
                log_func.logger.warning(f"Failed to clean up temporary directory: {tmp_index_dir}", e)


def load_existing_meta_data():
    """Loads existing meta data from _meta.json if it exists."""
    meta_file = os.path.join(config.get_output_dir(), "_meta.json")
    if os.path.exists(meta_file):
        with open(meta_file, "r", encoding="utf-8") as f:
            return json.load(f)
    return []


def get_articles_to_scrape(articles, existing_meta_data, update_mode=False, log_func=None):
    """Determines which articles need to be scraped based on content hashes and update mode."""
    if not update_mode or not existing_meta_data:
        if log_func:
            log_func.logger.info("Not in update mode or no existing data, scraping all articles")
        return articles  # If not in update mode or no existing data, scrape all
        
    # Create a lookup dictionary from existing metadata
    existing_lookup = {}
    for article in existing_meta_data:
        if 'url' in article:
            existing_lookup[article['url']] = article
    
    articles_to_scrape = []
    new_articles = 0
    modified_articles = 0
    unchanged_articles = 0
    
    for article in articles:
        url = article.get('url')
        if not url:
            articles_to_scrape.append(article)  # Include articles without URL
            continue
            
        existing_article = existing_lookup.get(url)
        if not existing_article:
            # New article that wasn't in previous scrape
            articles_to_scrape.append(article)
            new_articles += 1
        else:
            # Check if content hash is different
            existing_hash = existing_article.get('content_hash')
            current_hash = article.get('content_hash')
            
            if existing_hash != current_hash:
                # Content has changed, need to scrape
                articles_to_scrape.append(article)
                modified_articles += 1
            else:
                # Content unchanged, skip scraping
                unchanged_articles += 1
    
    if log_func:
        log_func.logger.info(f"Update mode: {new_articles} new articles, {modified_articles} modified articles, {unchanged_articles} unchanged articles")
        log_func.logger.info(f"Will scrape {len(articles_to_scrape)} articles out of {len(articles)} total")
    
    return articles_to_scrape


def should_force_reindex(toc_tree, existing_meta_data, log_func=None):
    """Checks if the TOC structure has changed significantly to warrant a force reindex."""
    if not existing_meta_data:
        if log_func:
            log_func.logger.info("No existing metadata found, proceeding with normal indexing")
        return False  # First time, don't force reindex since we want to process normally
        
    # Create mapping of URLs to articles from existing metadata
    existing_urls = set(article['url'] for article in existing_meta_data if 'url' in article)
    new_urls = set()
    
    # Helper to collect all URLs from the TOC tree
    def collect_urls(nodes):
        for node in nodes:
            if 'url' in node and node['url']:
                new_urls.add(node['url'])
            if 'children' in node and node['children']:
                collect_urls(node['children'])
    
    collect_urls(toc_tree)
    
    # Check if there are significant changes in URLs
    new_articles_count = len(new_urls - existing_urls)
    removed_articles_count = len(existing_urls - new_urls)
    
    # Log changes
    if log_func:
        log_func.logger.info(f"TOC changes: {new_articles_count} new articles, {removed_articles_count} removed articles")
        log_func.logger.info(f"Existing articles: {len(existing_urls)}, New articles: {len(new_urls)}")
    
    # If more than 20% of articles are new or removed, force reindex (download all)
    total_existing = len(existing_urls)
    if total_existing > 0:
        change_threshold = 0.2  # 20% threshold
        significant_change = (new_articles_count + removed_articles_count) / total_existing > change_threshold
        if significant_change and log_func:
            log_func.logger.warning(f"Significant TOC structure changes detected ({(new_articles_count + removed_articles_count) / total_existing * 100:.1f}%), forcing reindex")
        return significant_change
    
    return new_articles_count > 0  # If no existing, but have new ones, don't force since it's just adding

def _save_as_docx(soup, article_info, docx_file, log_func=None):
    """
    Saves article content as a DOCX file with proper formatting.
    
    Args:
        soup: BeautifulSoup object with article content
        article_info: Dictionary with article metadata
        docx_file: Path to save the DOCX file
        log_func: Logging function (optional)
    """
    if not DOCX_AVAILABLE:
        raise ImportError("python-docx library is not available")
    
    try:
        # Create a new Document
        doc = Document()
        
        # Add title
        title = article_info.get('title', 'Unknown Title')
        title_para = doc.add_heading(title, level=1)
        title_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # Add URL
        url = article_info.get('url', '')
        if url:
            url_para = doc.add_paragraph(f"URL: {url}")
            url_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # Add empty paragraph for spacing
        
        # Process HTML content
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'ul', 'ol', 'li', 'table', 'tr', 'td', 'th', 'blockquote']):
            if element.name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
                # Map HTML heading levels to DOCX heading levels
                level = min(int(element.name[1]), 9)  # DOCX supports up to level 9
                doc.add_heading(element.get_text(strip=True), level=level)
            
            elif element.name == 'p':
                text = element.get_text(strip=True)
                if text:
                    doc.add_paragraph(text)
            
            elif element.name == 'blockquote':
                text = element.get_text(strip=True)
                if text:
                    quote_para = doc.add_paragraph(text)
                    # Indent the quote
                    quote_para.paragraph_format.left_indent = Pt(36)
            
            elif element.name in ['ul', 'ol']:
                # Handle lists
                list_items = element.find_all('li', recursive=False)
                for li in list_items:
                    text = li.get_text(strip=True)
                    if text:
                        if element.name == 'ul':
                            doc.add_paragraph(text, style='List Bullet')
                        else:
                            doc.add_paragraph(text, style='List Number')
            
            elif element.name == 'table':
                # Handle tables
                rows = element.find_all('tr')
                if rows:
                    # Create table in DOCX
                    table = doc.add_table(rows=len(rows), cols=0)
                    table.style = 'Table Grid'
                    
                    for row_idx, row in enumerate(rows):
                        cells = row.find_all(['td', 'th'])
                        if row_idx == 0:  # First row, add columns
                            table.columns = len(cells)
                        
                        for col_idx, cell in enumerate(cells):
                            if col_idx < len(table.columns):
                                table.cell(row_idx, col_idx).text = cell.get_text(strip=True)
                                # Make header row bold
                                if row_idx == 0 and cell.name == 'th':
                                    table.cell(row_idx, col_idx).paragraphs[0].runs[0].bold = True
        
        # Save the document
        doc.save(docx_file)
        
        if log_func:
            log_func.logger.debug(f"Successfully created DOCX file: {docx_file}")
            
    except Exception as e:
        if log_func:
            log_func.logger.error(f"Error creating DOCX file: {str(e)}")
        raise
